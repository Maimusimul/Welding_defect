import os
import cv2
import docx

from docx import Document
from docx.shared import Inches

from PyQt5.QtWidgets import (
    QWidget, QLabel, QPushButton, QTextEdit, QVBoxLayout, QHBoxLayout, QGroupBox,
    QMainWindow, QMenuBar, QStatusBar, QFileDialog, QListWidget, QListWidgetItem,
    QComboBox, QSlider, QSplitter
)
from PyQt5.QtGui import QImage, QPixmap, QIcon
from PyQt5.QtCore import Qt, QSize
import numpy as np

from core.image_processor import ImageProcessor
from ui.image_viewer import ImageViewer




class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.loaded_image_paths = set()  # сохраняем пути уникальных изображений
        self.processor = ImageProcessor("models/YOLO_11.pt")
        self.setWindowTitle("Программа детектирования дефектов сварки")
        self.setGeometry(100, 100, 1200, 900)

        self.last_annotated = None

        # Центральный виджет и layout
        central_widget = QWidget()
        outer_layout = QVBoxLayout(central_widget)

        # Заголовок
        header = QLabel("<h2 align='center'>Программа детектирования дефектов сварки</h2>")
        header.setAlignment(Qt.AlignHCenter)
        outer_layout.addWidget(header)

        # ─────────── Блок с изображением ───────────
        self.label_image_name = QLabel("Файл не выбран")
        self.label_image_name.setAlignment(Qt.AlignCenter)
        self.label_image_name.setStyleSheet("font-weight: bold; font-size: 12pt;")

        self.label_imagePreview = QLabel("Здесь будет загруженное изображение")
        self.label_imagePreview.setAlignment(Qt.AlignCenter)
        self.label_imagePreview.setMinimumSize(400, 400)
        self.label_imagePreview.setStyleSheet("border: 1px solid gray")
        self.label_imagePreview.mousePressEvent = self.open_image_viewer

        preview_layout = QVBoxLayout()
        preview_layout.addWidget(self.label_image_name)
        preview_layout.addWidget(self.label_imagePreview)

        group_preview = QGroupBox("Предпросмотр")
        group_preview.setLayout(preview_layout)

        self.listWidget_images = QListWidget()
        self.listWidget_images.setMinimumHeight(120)
        self.listWidget_images.setIconSize(QSize(100, 100))
        self.listWidget_images.setGridSize(QSize(110, 110))  # фикс. размер ячеек, чтобы не прыгали
        self.listWidget_images.setViewMode(QListWidget.IconMode)
        self.listWidget_images.setResizeMode(QListWidget.Adjust)
        self.listWidget_images.itemClicked.connect(self.on_image_selected)

        self.button_load = QPushButton("Загрузить изображения")
        self.button_load.clicked.connect(self.open_images)

        gallery_layout = QVBoxLayout()
        gallery_layout.addWidget(self.listWidget_images)
        gallery_layout.addWidget(self.button_load)

        group_gallery = QGroupBox("Галерея изображений")
        group_gallery.setLayout(gallery_layout)

        # ─────────── Блок с информацией (теперь будет внизу) ───────────
        self.label_filename = QLabel("Файл не выбран")
        self.label_resultInfo = QLabel()
        self.label_resultInfo.setAlignment(Qt.AlignLeft | Qt.AlignTop)
        self.label_resultInfo.setWordWrap(True)

        self.comboBox_model = QComboBox()
        self.comboBox_model.addItems(["YOLO_11.pt", "YOLO_10.pt", "YOLO_8.pt"])
        self.comboBox_model.currentTextChanged.connect(self.on_model_changed)

        model_info_layout = QVBoxLayout()
        model_info_layout.addWidget(QLabel("Выбор модели"))
        model_info_layout.addWidget(self.comboBox_model)
        model_info_layout.addWidget(self.label_filename)
        model_info_layout.addWidget(self.label_resultInfo)

        group_model_info = QGroupBox("Модель и информация")
        group_model_info.setLayout(model_info_layout)

        # ─────────── Горизонтальный сплиттер: слева изображения, справа инфо ───────────
        top_splitter = QSplitter(Qt.Horizontal)
        top_splitter.addWidget(group_preview)  # Слева — превью
        top_splitter.addWidget(group_gallery)  # Справа — галерея

        # ─────────── Нижняя часть — QTextEdit ───────────
        self.textEdit = QTextEdit()
        group_text = QGroupBox("Информация о дефектах")
        text_layout = QVBoxLayout()
        text_layout.addWidget(self.textEdit)
        group_text.setLayout(text_layout)

        # ─────────── Вертикальный сплиттер: сверху top_splitter, снизу text ───────────
        vertical_splitter = QSplitter(Qt.Vertical)
        vertical_splitter.addWidget(top_splitter)
        vertical_splitter.addWidget(group_model_info)  # ⬅️ вставляем выбор модели под галереей и превью
        vertical_splitter.addWidget(group_text)

        outer_layout.addWidget(vertical_splitter)

        # ─────────── Кнопки снизу ───────────
        buttons_layout = QHBoxLayout()
        self.button_save = QPushButton("Сохранить результат")
        self.button_export = QPushButton("Экспортировать отчет")
        self.button_clear = QPushButton("Очистить интерфейс")
        self.button_exit = QPushButton("Завершить работу")

        self.button_export.clicked.connect(self.export_report)
        self.button_save.clicked.connect(self.save_result)
        self.button_clear.clicked.connect(self.clear_interface)
        self.button_exit.clicked.connect(self.close)

        for btn in [self.button_save, self.button_export, self.button_clear, self.button_exit]:
            buttons_layout.addWidget(btn)

        outer_layout.addLayout(buttons_layout)

        self.setCentralWidget(central_widget)
        self.setMenuBar(QMenuBar(self))
        self.setStatusBar(QStatusBar(self))

    def open_images(self):
        file_names, _ = QFileDialog.getOpenFileNames(
            self, "Выбрать изображения", "", "Изображения (*.png *.jpg *.jpeg *.bmp)"
        )
        if not file_names:
            return

        added = 0
        for path in file_names:
            if path in self.loaded_image_paths:
                continue  # Уже добавлен

            self.loaded_image_paths.add(path)
            item = QListWidgetItem()  # не указываем текст, чтобы не прыгали
            pixmap = QPixmap(path).scaled(100, 100, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            item.setIcon(QIcon(pixmap))
            item.setData(Qt.UserRole, path)
            item.setToolTip(os.path.basename(path))  # покажем имя при наведении
            self.listWidget_images.addItem(item)

            added += 1

        self.label_filename.setText(f"Загружено: {added} новых файл(ов)")

    def on_image_selected(self, item):
        path = item.data(Qt.UserRole)
        if not path or not os.path.exists(path):
            return

        self.label_filename.setText(os.path.basename(path))
        self.label_image_name.setText(os.path.basename(path))
        image = cv2.imread(path)

        try:
            annotated, detections = self.processor.detect_defects(image)  # ⬅️ теперь получаем два значения
            self.last_detections = detections  # ⬅️ сохраним для отчёта
            self.last_annotated = annotated
            num_defects = len(detections)
            if num_defects > 0:
                avg_conf = sum(d['confidence'] for d in detections) / num_defects
                max_conf = max(d['confidence'] for d in detections)
                stats_text = (
                    f"Обнаружено дефектов: {num_defects}\n"
                    f"Средняя уверенность: {avg_conf:.2f}\n"
                    f"Макс. уверенность: {max_conf:.2f}"
                )
            else:
                stats_text = "Дефекты не обнаружены."
            self.show_image_in_label(annotated)  # ⬅️ ВАЖНО: отобразить превью
            self.label_resultInfo.setText("Обработка завершена.")
            # Заполняем поле "Информация о дефектах"
            self.textEdit.setPlainText(
                f"Файл: {os.path.basename(path)}\n"
                f"Использована модель: {self.comboBox_model.currentText()}\n"
                f"{stats_text}"
            )


        except Exception as e:
            self.label_resultInfo.setText(f"Ошибка: {e}")

    def apply_zoom(self):
        if self.last_annotated is None:
            return
        zoom = self.current_zoom / 100
        img = self.last_annotated
        rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        h, w, ch = rgb.shape
        new_w, new_h = int(w * zoom), int(h * zoom)
        resized = cv2.resize(rgb, (new_w, new_h), interpolation=cv2.INTER_LINEAR)
        qimg = QImage(resized.data, new_w, new_h, new_w * ch, QImage.Format_RGB888)
        pixmap = QPixmap.fromImage(qimg)
        scaled = pixmap.scaled(self.label_imagePreview.size(), Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.label_imagePreview.setPixmap(scaled)

    def update_zoom(self, value):
        self.current_zoom = value
        self.apply_zoom()

    def open_image_viewer(self, event):
        if self.last_annotated is not None:
            viewer = ImageViewer(self.last_annotated, self)
            viewer.exec_()

    def on_model_changed(self, model_name):
        self.processor = ImageProcessor(os.path.join("models", model_name))
        self.label_resultInfo.setText(f"Модель переключена на: {model_name}")

    def show_image_in_label(self, image_bgr):
        image_rgb = cv2.cvtColor(image_bgr, cv2.COLOR_BGR2RGB)
        h, w, ch = image_rgb.shape
        bytes_per_line = ch * w
        qimg = QImage(image_rgb.data, w, h, bytes_per_line, QImage.Format_RGB888)
        pixmap = QPixmap.fromImage(qimg)
        scaled = pixmap.scaled(self.label_imagePreview.size(), Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.label_imagePreview.setPixmap(scaled)

    def export_report(self):
        if self.last_annotated is None:
            self.label_resultInfo.setText("Нет изображения для экспорта.")
            return

        doc = Document()

        # Устанавливаем общий стиль документа
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = docx.shared.Pt(14)

        # Заголовок отчета
        heading = doc.add_heading("Отчёт по анализу сварочного изображения", level=1)
        heading.style.font.name = 'Times New Roman'
        heading.style.font.size = docx.shared.Pt(14)
        heading.style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

        # Информация о дефектах
        doc.add_heading("Описание дефектов", level=1)
        doc.add_paragraph(self.textEdit.toPlainText())

        # Обнаруженные объекты
        if hasattr(self, 'last_detections') and self.last_detections:
            doc.add_heading("Обнаруженные объекты", level=1)
            for det in self.last_detections:
                label = det["label"]
                conf = det["confidence"]
                doc.add_paragraph(f"{label} — {conf:.2f}", style='List Number')
        else:
            doc.add_paragraph("Объекты не обнаружены.")

        # Сохраняем аннотированное изображение во временный файл
        annotated_path = "temp_annotated.png"
        cv2.imwrite(annotated_path, self.last_annotated)
        doc.add_heading("Аннотированное изображение", level=1)
        doc.add_picture(annotated_path, width=Inches(5))

        # Сохранить как
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить отчёт",
            "report.docx",
            "Документ Word (*.docx)"
        )
        if save_path:
            doc.save(save_path)
            self.label_resultInfo.setText(f"Отчёт сохранён: {save_path}")
        else:
            self.label_resultInfo.setText("Сохранение отменено.")

        # Удаляем временное изображение
        if os.path.exists(annotated_path):
            os.remove(annotated_path)

    # Сохранение результата
    def save_result(self):
        if self.last_annotated is None:
            self.label_resultInfo.setText("Нет результата для сохранения.")
            return
        path, _ = QFileDialog.getSaveFileName(self, "Сохранить изображение", "", "PNG (*.png);;JPEG (*.jpeg);;JPG (*.jpg);;BMP (*.bmp)")
        if path:
            cv2.imwrite(path, self.last_annotated)
            self.label_resultInfo.setText(f"Изображение сохранено: {path}")
        else:
            self.label_resultInfo.setText("Сохранение отменено.")

    # Очистка интерфейса
    def clear_interface(self):
        self.label_imagePreview.clear()
        self.label_filename.setText("Файл не выбран")
        self.label_image_name.setText("Файл не выбран")
        self.label_resultInfo.clear()
        self.loaded_image_paths.clear()
        self.textEdit.clear()
        self.last_annotated = None
